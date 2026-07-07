"""
Generate a Word (.docx) and PDF of the Bring Back the Salmon one-pager content.
No images; formatted for colleague editing.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, KeepTogether
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

# ── Colour palette ──────────────────────────────────────────────────────────
NAVY   = (0x0b/255, 0x3a/255, 0x53/255)   # #0b3a53
FOREST = (0x2e/255, 0x5b/255, 0x3a/255)   # #2e5b3a
ACCENT = (0xb7/255, 0x65/255, 0x1f/255)   # #b7651f
INK    = (0x0f/255, 0x1d/255, 0x24/255)   # #0f1d24
MUTE   = (0x60/255, 0x70/255, 0x78/255)   # #607078

RL_NAVY   = colors.Color(*NAVY)
RL_FOREST = colors.Color(*FOREST)
RL_ACCENT = colors.Color(*ACCENT)
RL_INK    = colors.Color(*INK)
RL_MUTE   = colors.Color(*MUTE)
RL_SAND   = colors.Color(0xed/255, 0xe6/255, 0xd2/255)
RL_WHITE  = colors.white
RL_LINE   = colors.Color(0xd9/255, 0xd3/255, 0xc3/255)

# ── Shared content ───────────────────────────────────────────────────────────

PROGRAM_NAME = "Bring Back the Atlantic Salmon"
SUBTITLE      = "Lake Ontario Restoration Program"
PARTNERS      = "Ontario Ministry of Natural Resources · Ontario Federation of Anglers and Hunters"
WEBSITE       = "bringbackthesalmon.ca"
CONTACT       = "info@bringbackthesalmon.ca"

HERO_KICKER = "The return of a Lake Ontario icon"
HERO_H1     = "Restoring Atlantic Salmon to Lake Ontario"

OUR_GOAL_HEAD  = "Restore Atlantic Salmon and support a diverse salmon and trout fishery."
OUR_GOAL_BODY  = (
    "This program is working to restore Atlantic Salmon to Lake Ontario while creating "
    "a unique recreational fishery that benefits Ontario's economy. At the same time, "
    "it improves critical habitat that strengthens biodiversity and honours the lake's "
    "cultural heritage."
)
OUR_GOAL_INVESTING = (
    "OFAH Classroom Hatchery Program inspires the next generation by connecting "
    "thousands of students to hands-on learning in salmon biology, conservation, "
    "and stewardship."
)

WHY_HEAD = "Better habitat. Better fishing. A better lake."
WHY_BODY = (
    "Atlantic Salmon are returning to the fishery in record numbers. Guided by years "
    "of experience, the program uses an adaptive management approach that focuses on "
    "targeted stocking of older life stages of fish, increasing adult abundance, "
    "tributary habitat restoration and enhancement. These efforts benefit fish and "
    "wildlife across the Lake Ontario watershed while supporting outstanding fishing "
    "opportunities, from local shorelines and tributaries to a thriving offshore boat fishery."
)

PRIORITIES = [
    ("01", "Increase adult Atlantic Salmon abundance in Lake Ontario"),
    ("02", "Continue to improve habitat in priority watersheds"),
    ("03", "Focus on stocking older life stage fish directly into Lake Ontario"),
    ("04", "Enhance Atlantic Salmon recreational angling opportunities"),
    ("05", "Improve communication and awareness"),
]

LAKE_PARAS = [
    (
        "Fish do not recognize borders, so through the program's adaptive management "
        "approach it collaborates with binational agency partners in New York State, "
        "on Lake wide stocking initiatives, habitat work, and science to support "
        "Atlantic Salmon restoration."
    ),
    (
        "Biologists are studying returning adult Atlantic Salmon to guide future "
        "restoration efforts. Sampling more than 300 adults over the past three years "
        "has provided an unprecedented opportunity to better understand this population "
        "and help advance its recovery."
    ),
    (
        "Catch rates are climbing, with anglers reporting more Atlantic Salmon than "
        "seen in decades. Lake-wide restoration efforts are supporting a strong and "
        "diverse salmon and trout fishery. Anglers across Lake Ontario can now enjoy "
        "expanded fishing opportunities, from accessible nearshore and tributary waters "
        "to a thriving offshore fishery."
    ),
]

LAKE_LEADIN = (
    "Get out on the water, experience the fishery and report your Atlantic Salmon "
    "catches to help support ongoing restoration efforts."
)

REPORT_CTA_HEAD = "Help Track the Return"
REPORT_CTA_BODY = (
    f"Have you caught or seen an Atlantic Salmon? Report your catch to "
    f"{CONTACT} or visit {WEBSITE}."
)

# Tuples: (number_or_None, label, description)
# None = text-only KPI (no big number)
KPIS = [
    (None,    "Increased Catch Rates across Lake Ontario",
              "Anglers across the lake are reporting Atlantic Salmon catches at increasing rates."),
    ("300+",  "Increased Adult Returns",
              "Returning Atlantic Salmon captured and studied in the past three years."),
    ("150+",  "Habitat projects",
              "Stream and watershed improvements delivered by OFAH and partners."),
    ("1,500+","Students reached",
              "Through the OFAH Classroom Hatchery & outreach program."),
    ("50+",   "Partner organizations",
              "Working together on fish production, habitat restoration, and outreach."),
]

DOING = [
    ("Fish Production",
     "Stocking Atlantic Salmon at older life stages where survival, returns, and restoration impact come together."),
    ("Habitat Restoration",
     "Improving streams, shorelines, and watersheds. These efforts benefit fish and wildlife across Lake Ontario."),
    ("Science & Monitoring",
     "Tracking fish using tags, cameras, and angler surveys to guide better decisions."),
    ("Education & Outreach",
     "Connect with anglers, students, and communities through OFAH programs across Ontario."),
]


# ═══════════════════════════════════════════════════════════════════════════════
# WORD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════

def rgb(r, g, b):
    return RGBColor(r, g, b)

NAVY_RGB   = rgb(0x0b, 0x3a, 0x53)
FOREST_RGB = rgb(0x2e, 0x5b, 0x3a)
ACCENT_RGB = rgb(0xb7, 0x65, 0x1f)
INK_RGB    = rgb(0x0f, 0x1d, 0x24)
MUTE_RGB   = rgb(0x60, 0x70, 0x78)

def add_heading(doc, text, level=1, color=None, size=None, bold=True, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size or (18 if level == 1 else 14 if level == 2 else 12))
    if color:
        run.font.color.rgb = color
    return p

def add_body(doc, text, size=10.5, color=None, space_before=2, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_kicker(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = color or FOREST_RGB
    return p

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def build_docx(out_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    # ── MASTHEAD ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(PROGRAM_NAME.upper())
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY_RGB

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(12)
    r2 = p2.add_run(SUBTITLE)
    r2.font.size = Pt(10)
    r2.font.color.rgb = MUTE_RGB

    # Divider
    hr = doc.add_paragraph()
    pPr = hr._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),  'single')
    bottom.set(qn('w:sz'),   '12')
    bottom.set(qn('w:space'),'1')
    bottom.set(qn('w:color'),'0b3a53')
    pBdr.append(bottom)
    pPr.append(pBdr)
    hr.paragraph_format.space_after = Pt(8)

    # ── HERO ──────────────────────────────────────────────────────────────────
    add_body(doc, HERO_KICKER.upper(), size=8, color=MUTE_RGB, space_before=0, space_after=2)
    add_heading(doc, HERO_H1, level=1, color=NAVY_RGB, size=20, space_before=0, space_after=14)

    # ── OUR GOAL / WHY THIS MATTERS ──────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.autofit = False
    tbl.columns[0].width = Inches(3.25)
    tbl.columns[1].width = Inches(3.25)

    for idx, (kicker, head, body) in enumerate([
        ("Our Goal",        OUR_GOAL_HEAD, OUR_GOAL_BODY),
        ("Why This Matters", WHY_HEAD,     WHY_BODY),
    ]):
        cell = tbl.rows[0].cells[idx]
        cell.width = Inches(3.25)

        # remove table borders for cleaner look
        kp = cell.add_paragraph()
        kp.paragraph_format.space_before = Pt(4)
        kp.paragraph_format.space_after  = Pt(4)
        kr = kp.add_run(kicker.upper())
        kr.bold = True
        kr.font.size = Pt(8)
        kr.font.color.rgb = FOREST_RGB

        hp = cell.add_paragraph()
        hp.paragraph_format.space_after = Pt(6)
        hr2 = hp.add_run(head)
        hr2.bold = True
        hr2.font.size = Pt(13)
        hr2.font.color.rgb = INK_RGB

        bp = cell.add_paragraph(body)
        bp.runs[0].font.size = Pt(9.5)
        if idx == 0:
            inv_p = cell.add_paragraph()
            inv_p.paragraph_format.space_before = Pt(4)
            inv_p.paragraph_format.space_after  = Pt(2)
            inv_bold = inv_p.add_run("Investing in the Future: ")
            inv_bold.bold = True
            inv_bold.font.size = Pt(9.5)
            inv_bold.font.color.rgb = INK_RGB
            inv_rest = inv_p.add_run(OUR_GOAL_INVESTING)
            inv_rest.font.size = Pt(9.5)

    doc.add_paragraph()  # spacer

    # ── FIVE PRIORITIES ───────────────────────────────────────────────────────
    add_heading(doc, "Ontario's Five Priorities", level=2, color=INK_RGB, size=14, space_before=14)

    ptbl = doc.add_table(rows=1, cols=5)
    ptbl.style = 'Table Grid'
    ptbl.autofit = False
    col_w = Inches(1.3)
    for i in range(5):
        ptbl.columns[i].width = col_w

    row = ptbl.rows[0]
    for i, (num, text) in enumerate(PRIORITIES):
        cell = row.cells[i]
        cell.width = col_w

        np = cell.add_paragraph()
        np.paragraph_format.space_before = Pt(4)
        np.paragraph_format.space_after  = Pt(2)
        nr = np.add_run(num)
        nr.bold = True
        nr.font.size = Pt(16)
        nr.font.color.rgb = FOREST_RGB

        tp = cell.add_paragraph(text)
        tp.runs[0].font.size = Pt(8.5)
        tp.runs[0].bold = True
        tp.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # ── LAKE-WIDE APPROACH ────────────────────────────────────────────────────
    add_heading(doc, "A Lake-Wide Approach", level=2, color=INK_RGB, size=14, space_before=14)

    ltbl = doc.add_table(rows=1, cols=2)
    ltbl.style = 'Table Grid'
    ltbl.autofit = False
    ltbl.columns[0].width = Inches(3.8)
    ltbl.columns[1].width = Inches(2.7)

    # Left copy
    lcell = ltbl.rows[0].cells[0]
    lcell.width = Inches(3.8)
    for para_text in LAKE_PARAS:
        p = lcell.add_paragraph(para_text)
        p.runs[0].font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(6)

    # "Be part of the return!" lead-in
    leadin_p = lcell.add_paragraph()
    leadin_p.paragraph_format.space_before = Pt(2)
    leadin_p.paragraph_format.space_after  = Pt(6)
    lb = leadin_p.add_run("Be part of the return! ")
    lb.bold = True
    lb.font.size = Pt(9.5)
    lr = leadin_p.add_run(LAKE_LEADIN)
    lr.font.size = Pt(9.5)

    # Report CTA
    ctap = lcell.add_paragraph()
    ctap.paragraph_format.space_before = Pt(6)
    r_bold = ctap.add_run(REPORT_CTA_HEAD + ": ")
    r_bold.bold = True
    r_bold.font.size = Pt(9.5)
    r_bold.font.color.rgb = ACCENT_RGB
    r_rest = ctap.add_run(REPORT_CTA_BODY)
    r_rest.font.size = Pt(9.5)

    # Right sidebar (KPIs)
    rcell = ltbl.rows[0].cells[1]
    rcell.width = Inches(2.7)
    set_cell_bg(rcell, 'EDE6D2')

    sh = rcell.add_paragraph("Progress at a Glance")
    sh.runs[0].bold = True
    sh.runs[0].font.size = Pt(11)
    sh.paragraph_format.space_before = Pt(6)
    sh.paragraph_format.space_after  = Pt(8)

    for num, label, desc in KPIS:
        kp = rcell.add_paragraph()
        kp.paragraph_format.space_before = Pt(4)
        kp.paragraph_format.space_after  = Pt(1)
        if num is not None:
            nr = kp.add_run(num + "  ")
            nr.bold = True
            nr.font.size = Pt(16)
            nr.font.color.rgb = NAVY_RGB
        lr = kp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(11) if num is None else Pt(8.5)
        lr.font.color.rgb = NAVY_RGB if num is None else INK_RGB

        dp = rcell.add_paragraph(desc)
        dp.runs[0].font.size = Pt(8)
        dp.runs[0].font.color.rgb = MUTE_RGB
        dp.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    # ── WHAT WE'RE DOING ──────────────────────────────────────────────────────
    add_heading(doc, "What We're Doing on the Water", level=2, color=INK_RGB, size=14, space_before=14)

    dtbl = doc.add_table(rows=1, cols=4)
    dtbl.style = 'Table Grid'
    dtbl.autofit = False
    dw = Inches(1.625)
    for i in range(4):
        dtbl.columns[i].width = dw

    drow = dtbl.rows[0]
    for i, (title, desc) in enumerate(DOING):
        cell = drow.cells[i]
        cell.width = dw
        set_cell_bg(cell, '0b3a53')

        tp = cell.add_paragraph(title.upper())
        tp.runs[0].bold = True
        tp.runs[0].font.size = Pt(8.5)
        tp.runs[0].font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after  = Pt(4)

    # Second row for descriptions (white bg)
    dtbl.add_row()
    for i, (title, desc) in enumerate(DOING):
        cell = dtbl.rows[1].cells[i]
        cell.width = dw
        dp = cell.add_paragraph(desc)
        dp.runs[0].font.size = Pt(8.5)
        dp.paragraph_format.space_before = Pt(4)
        dp.paragraph_format.space_after  = Pt(6)

    doc.add_paragraph()

    # ── FOOTER ────────────────────────────────────────────────────────────────
    hr2 = doc.add_paragraph()
    pPr2 = hr2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    top2 = OxmlElement('w:top')
    top2.set(qn('w:val'),  'single')
    top2.set(qn('w:sz'),   '12')
    top2.set(qn('w:space'),'1')
    top2.set(qn('w:color'),'0b3a53')
    pBdr2.append(top2)
    pPr2.append(pBdr2)
    hr2.paragraph_format.space_before = Pt(14)
    hr2.paragraph_format.space_after  = Pt(4)

    fp = doc.add_paragraph()
    fp.paragraph_format.space_before = Pt(2)
    fp.paragraph_format.space_after  = Pt(2)
    fb = fp.add_run("Delivered in Partnership: ")
    fb.bold = True
    fb.font.size = Pt(9)
    fb.font.color.rgb = NAVY_RGB
    fr = fp.add_run(PARTNERS)
    fr.font.size = Pt(9)

    wp = doc.add_paragraph()
    wp.paragraph_format.space_before = Pt(2)
    wr = wp.add_run(f"Learn more. Get involved.  |  {WEBSITE}")
    wr.bold = True
    wr.font.size = Pt(10)
    wr.font.color.rgb = NAVY_RGB

    doc.save(out_path)
    print(f"Word document saved: {out_path}")


# ═══════════════════════════════════════════════════════════════════════════════
# PDF
# ═══════════════════════════════════════════════════════════════════════════════

def build_pdf(out_path):
    doc = SimpleDocTemplate(
        out_path,
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch,
    )

    styles = getSampleStyleSheet()

    def S(name, parent='Normal', **kw):
        return ParagraphStyle(name, parent=styles[parent], **kw)

    sKicker = S('Kicker', fontSize=8,  textColor=RL_FOREST, leading=10, spaceBefore=10, spaceAfter=2,
                fontName='Helvetica-Bold')
    sH1     = S('H1',     fontSize=20, textColor=RL_NAVY,   leading=24, spaceBefore=4, spaceAfter=10,
                fontName='Helvetica-Bold')
    sH2     = S('H2',     fontSize=13, textColor=RL_INK,    leading=16, spaceBefore=10, spaceAfter=4,
                fontName='Helvetica-Bold')
    sH3     = S('H3',     fontSize=11, textColor=RL_INK,    leading=14, spaceBefore=8, spaceAfter=3,
                fontName='Helvetica-Bold')
    sBody   = S('Body',   fontSize=9.5,textColor=RL_INK,    leading=14, spaceBefore=2, spaceAfter=4,
                fontName='Helvetica')
    sSmall  = S('Small',  fontSize=8,  textColor=RL_MUTE,   leading=11, spaceBefore=0, spaceAfter=4,
                fontName='Helvetica')
    sMono   = S('Mono',   fontSize=8,  textColor=RL_INK,    leading=11, spaceBefore=0, spaceAfter=2,
                fontName='Helvetica-Bold')
    sNavyBold = S('NavyBold', fontSize=9, textColor=RL_NAVY, leading=12, spaceBefore=2, spaceAfter=2,
                  fontName='Helvetica-Bold')
    sAccent   = S('Accent',   fontSize=9.5, textColor=RL_ACCENT, leading=13, spaceBefore=6, spaceAfter=4,
                  fontName='Helvetica-Bold')
    sCenterWhite = S('CenterWhite', fontSize=8.5, textColor=RL_WHITE, leading=11, spaceBefore=4, spaceAfter=4,
                     fontName='Helvetica-Bold', alignment=TA_CENTER)

    story = []

    # ── MASTHEAD ──
    story.append(Paragraph(PROGRAM_NAME.upper(), S('M', fontSize=14, textColor=RL_NAVY,
                                                    fontName='Helvetica-Bold', leading=17, spaceBefore=0, spaceAfter=2)))
    story.append(Paragraph(SUBTITLE, S('MS', fontSize=10, textColor=RL_MUTE,
                                        fontName='Helvetica', leading=13, spaceBefore=0, spaceAfter=6)))
    story.append(HRFlowable(width='100%', thickness=3, color=RL_NAVY, spaceAfter=10))

    # ── HERO ──
    story.append(Paragraph(HERO_KICKER.upper(), sKicker))
    story.append(Paragraph(HERO_H1, sH1))

    # ── OUR GOAL / WHY THIS MATTERS ──
    w = (doc.width - 12) / 2
    intro_data = [[
        [Paragraph("OUR GOAL", sKicker),
         Paragraph(OUR_GOAL_HEAD, sH3),
         Paragraph(OUR_GOAL_BODY, sBody),
         Paragraph('<b>Investing in the Future:</b> ' + OUR_GOAL_INVESTING, sBody)],
        [Paragraph("WHY THIS MATTERS", sKicker),
         Paragraph(WHY_HEAD, sH3),
         Paragraph(WHY_BODY, sBody)],
    ]]
    intro_tbl = Table(intro_data, colWidths=[w, w])
    intro_tbl.setStyle(TableStyle([
        ('VALIGN',      (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 6),
        ('RIGHTPADDING',(0,0), (-1,-1), 6),
        ('LINEBELOW',   (0,0), (-1,-1), 0.5, RL_LINE),
    ]))
    story.append(intro_tbl)
    story.append(Spacer(1, 10))

    # ── FIVE PRIORITIES ──
    story.append(HRFlowable(width='100%', thickness=1, color=RL_LINE, spaceBefore=4, spaceAfter=6))
    story.append(Paragraph("Ontario's Five Priorities", sH2))

    cw = doc.width / 5
    pdata = [[
        [Paragraph(f'<font color="#2e5b3a"><b>{num}</b></font>', S('PN',fontSize=16,leading=18,fontName='Helvetica-Bold',spaceBefore=4,spaceAfter=2)),
         Paragraph(text, S('PT',fontSize=8,leading=11,fontName='Helvetica-Bold',spaceAfter=4))]
        for num, text in PRIORITIES
    ]]
    ptbl = Table(pdata, colWidths=[cw]*5)
    ptbl.setStyle(TableStyle([
        ('VALIGN',       (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING',  (0,0), (-1,-1), 6),
        ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ('BOX',          (0,0), (-1,-1), 0.5, RL_LINE),
        ('LINEBEFORE',   (1,0), (-1,-1), 0.5, RL_LINE),
    ]))
    story.append(ptbl)
    story.append(Spacer(1, 10))

    # ── LAKE-WIDE APPROACH ──
    story.append(HRFlowable(width='100%', thickness=1, color=RL_LINE, spaceBefore=4, spaceAfter=6))
    story.append(Paragraph("A Lake-Wide Approach", sH2))

    lw_copy  = doc.width * 0.58
    lw_side  = doc.width * 0.42 - 10

    copy_cells = [Paragraph(p, sBody) for p in LAKE_PARAS]
    copy_cells.append(Paragraph(
        '<b>Be part of the return!</b> ' + LAKE_LEADIN, sBody
    ))
    copy_cells.append(Spacer(1, 4))
    copy_cells.append(Paragraph(
        f'<font color="#b7651f"><b>{REPORT_CTA_HEAD}:</b></font> '
        f'{REPORT_CTA_BODY.replace(CONTACT, f"<font color=#14506e>{CONTACT}</font>").replace(WEBSITE, f"<font color=#14506e>{WEBSITE}</font>")}',
        sBody
    ))

    kpi_cells = [Paragraph("Progress at a Glance", sH3)]
    for num, label, desc in KPIS:
        kpi_cells.append(Spacer(1, 3))
        if num is not None:
            kpi_cells.append(Paragraph(
                f'<font color="#0b3a53" size="15"><b>{num}</b></font>  '
                f'<font size="8.5"><b>{label}</b></font>',
                S('KH', fontSize=9, leading=14, fontName='Helvetica', spaceBefore=2, spaceAfter=1)
            ))
        else:
            kpi_cells.append(Paragraph(
                f'<font color="#0b3a53"><b>{label}</b></font>',
                S('KH2', fontSize=10.5, leading=13, fontName='Helvetica-Bold', spaceBefore=2, spaceAfter=1)
            ))
        kpi_cells.append(Paragraph(desc, sSmall))

    lake_data = [[copy_cells, kpi_cells]]
    lake_tbl  = Table(lake_data, colWidths=[lw_copy, lw_side])
    lake_tbl.setStyle(TableStyle([
        ('VALIGN',       (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING',  (0,0), (-1,-1), 0),
        ('RIGHTPADDING', (0,0), (-1,-1), 0),
        ('LEFTPADDING',  (1,0), (1,-1),  8),
        ('BACKGROUND',   (1,0), (1,-1),  RL_SAND),
        ('LINEABOVE',    (1,0), (1,0),   2.5, RL_FOREST),
        ('TOPPADDING',   (1,0), (1,-1),  8),
        ('BOTTOMPADDING',(1,0), (1,-1),  8),
    ]))
    story.append(lake_tbl)
    story.append(Spacer(1, 10))

    # ── WHAT WE'RE DOING ──
    story.append(HRFlowable(width='100%', thickness=1, color=RL_LINE, spaceBefore=4, spaceAfter=6))
    story.append(Paragraph("What We're Doing on the Water", sH2))

    dw = doc.width / 4
    do_data = [
        [Paragraph(t.upper(), sCenterWhite) for t, _ in DOING],
        [Paragraph(d, S('Dd', fontSize=8.5, leading=12, fontName='Helvetica',
                         spaceBefore=4, spaceAfter=6, alignment=TA_LEFT)) for _, d in DOING],
    ]
    do_tbl = Table(do_data, colWidths=[dw]*4)
    do_tbl.setStyle(TableStyle([
        ('BACKGROUND',   (0,0), (-1,0),  RL_NAVY),
        ('TEXTCOLOR',    (0,0), (-1,0),  RL_WHITE),
        ('VALIGN',       (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING',  (0,0), (-1,-1), 8),
        ('RIGHTPADDING', (0,0), (-1,-1), 8),
        ('TOPPADDING',   (0,0), (-1,0),  6),
        ('BOTTOMPADDING',(0,0), (-1,0),  6),
        ('BOX',          (0,0), (-1,-1), 0.5, RL_LINE),
        ('LINEBEFORE',   (1,0), (-1,-1), 0.5, RL_LINE),
    ]))
    story.append(do_tbl)
    story.append(Spacer(1, 10))

    # ── FOOTER ──
    story.append(HRFlowable(width='100%', thickness=3, color=RL_NAVY, spaceBefore=8, spaceAfter=6))
    story.append(Paragraph(
        f'<b>Delivered in Partnership:</b> {PARTNERS}',
        S('Foot', fontSize=9, textColor=RL_INK, fontName='Helvetica', leading=13, spaceBefore=2, spaceAfter=2)
    ))
    story.append(Paragraph(
        f'Learn more. Get involved.  |  <font color="#0b3a53"><b>{WEBSITE}</b></font>',
        S('FootCTA', fontSize=10, textColor=RL_INK, fontName='Helvetica', leading=13, spaceBefore=2, spaceAfter=0)
    ))

    doc.build(story)
    print(f"PDF saved: {out_path}")


if __name__ == '__main__':
    base = r"C:\Users\lomu-\Documents\GitHub\lake-ontario-atlantic-salmon-communication-plan\project\Lake Ontario Atlantic Salmon Communication Plan"
    build_docx(f"{base}\\Bring Back the Salmon - Program Brief.docx")
    build_pdf(f"{base}\\Bring Back the Salmon - Program Brief.pdf")
